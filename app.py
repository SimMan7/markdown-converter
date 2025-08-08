import os
import logging
import markdown
import requests
import json
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, Response, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import uuid
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Create Flask app
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "dev-secret-key-change-in-production")

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'md', 'markdown'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Create upload directory if it doesn't exist
try:
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
except Exception as e:
    logging.error(f"Error creating upload folder: {e}")

# In-memory storage for Vercel (since filesystem is read-only)
file_storage = {}

# reCAPTCHA configuration
RECAPTCHA_SECRET_KEY = "6LeSYJ0rAAAAABx3xOqWudqBdr36gK6IcTUnBgaK"
RECAPTCHA_VERIFY_URL = "https://www.google.com/recaptcha/api/siteverify"

def allowed_file(filename):
    """Check if file has allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_old_files():
    """Remove files older than 1 hour from uploads directory"""
    try:
        current_time = datetime.now()
        for filename in os.listdir(UPLOAD_FOLDER):
            if filename == '.gitkeep':
                continue
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            if os.path.isfile(filepath):
                file_time = datetime.fromtimestamp(os.path.getctime(filepath))
                if (current_time - file_time).seconds > 3600:  # 1 hour
                    os.remove(filepath)
                    logging.info(f"Removed old file: {filename}")
    except Exception as e:
        logging.error(f"Error cleaning up files: {e}")

def verify_recaptcha(recaptcha_response):
    """Verify reCAPTCHA response with Google's API"""
    try:
        data = {
            'secret': RECAPTCHA_SECRET_KEY,
            'response': recaptcha_response
        }
        
        response = requests.post(RECAPTCHA_VERIFY_URL, data=data, timeout=10)
        result = response.json()
        
        logging.info(f"reCAPTCHA verification result: {result}")
        
        return result.get('success', False)
    except Exception as e:
        logging.error(f"Error verifying reCAPTCHA: {e}")
        return False

@app.route('/')
def index():
    """Main page"""
    try:
        cleanup_old_files()
        return render_template('index.html')
    except Exception as e:
        logging.error(f"Error in index route: {e}")
        return render_template('index.html')

@app.route('/ads.txt')
def ads_txt():
    """Serve ads.txt file for Google AdSense"""
    try:
        return Response(
            "google.com, ca-pub-6865885814212781, DIRECT, f08c47fec0942fa0",
            mimetype='text/plain',
            headers={'Cache-Control': 'public, max-age=3600'}
        )
    except Exception as e:
        logging.error(f"Error serving ads.txt: {e}")
        return Response("google.com, ca-pub-6865885814212781, DIRECT, f08c47fec0942fa0", mimetype='text/plain')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload"""
    try:
        if 'file' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('index'))
        
        file = request.files['file']
        
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('index'))
        
        if not allowed_file(file.filename):
            flash('Please upload a valid Markdown file (.md or .markdown)', 'error')
            return redirect(url_for('index'))
        
        # Generate unique filename to avoid conflicts
        unique_id = str(uuid.uuid4())
        filename = f"{unique_id}_{secure_filename(file.filename or 'upload')}"
        
        # Read file content
        content = file.read().decode('utf-8')
        
        # Store in memory (in production, use cloud storage)
        file_storage[filename] = {
            'content': content,
            'original_filename': file.filename,
            'upload_time': datetime.now()
        }
        
        # Convert markdown to HTML for preview
        html_content = markdown.markdown(
            content, 
            extensions=['extra', 'tables', 'codehilite', 'toc']
        )
        
        flash('File uploaded successfully!', 'success')
        return render_template('index.html', 
                             preview_content=html_content, 
                             filename=filename,
                             original_filename=file.filename)
        
    except Exception as e:
        logging.error(f"Upload error: {e}")
        flash(f'Error uploading file: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/download/<format>/<filename>')
def download_file(format, filename):
    """Download converted file"""
    try:
        if filename not in file_storage:
            flash('File not found. Please upload a file first.', 'error')
            return redirect(url_for('index'))
        
        file_data = file_storage[filename]
        md_content = file_data['content']
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content, 
            extensions=['extra', 'tables', 'codehilite', 'toc']
        )
        
        if format == 'pdf':
            # Create beautifully formatted HTML for PDF conversion
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Markdown Document</title>
                <style>
                    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
                    
                    * {{
                        margin: 0;
                        padding: 0;
                        box-sizing: border-box;
                    }}
                    
                    body {{
                        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                        line-height: 1.7;
                        color: #1a202c;
                        max-width: 800px;
                        margin: 0 auto;
                        padding: 40px 20px;
                        background-color: white;
                        font-size: 16px;
                    }}
                    
                    h1, h2, h3, h4, h5, h6 {{
                        color: #2d3748;
                        margin-top: 40px;
                        margin-bottom: 20px;
                        page-break-after: avoid;
                        font-weight: 600;
                        line-height: 1.3;
                    }}
                    
                    h1 {{
                        font-size: 2.5rem;
                        border-bottom: 3px solid #667eea;
                        padding-bottom: 15px;
                        margin-top: 0;
                        margin-bottom: 30px;
                        color: #1a202c;
                    }}
                    
                    h2 {{
                        font-size: 2rem;
                        border-bottom: 2px solid #e2e8f0;
                        padding-bottom: 10px;
                        margin-top: 35px;
                        color: #2d3748;
                    }}
                    
                    h3 {{
                        font-size: 1.5rem;
                        color: #4a5568;
                        margin-top: 30px;
                    }}
                    
                    h4 {{
                        font-size: 1.25rem;
                        color: #4a5568;
                        margin-top: 25px;
                    }}
                    
                    p {{
                        margin-bottom: 1.5rem;
                        text-align: justify;
                        color: #2d3748;
                    }}
                    
                    code {{
                        background-color: #f7fafc;
                        color: #e53e3e;
                        padding: 4px 8px;
                        border-radius: 6px;
                        font-family: 'Monaco', 'Courier New', 'Fira Code', monospace;
                        font-size: 0.9em;
                        border: 1px solid #e2e8f0;
                    }}
                    
                    pre {{
                        background-color: #f7fafc;
                        border: 1px solid #e2e8f0;
                        border-radius: 8px;
                        padding: 20px;
                        overflow-x: auto;
                        margin: 25px 0;
                        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                    }}
                    
                    pre code {{
                        background-color: transparent;
                        padding: 0;
                        color: #2d3748;
                        border: none;
                        font-size: 0.95em;
                        line-height: 1.6;
                    }}
                    
                    blockquote {{
                        border-left: 4px solid #667eea;
                        margin: 25px 0;
                        padding: 15px 25px;
                        color: #4a5568;
                        font-style: italic;
                        background-color: #f7fafc;
                        border-radius: 0 8px 8px 0;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
                    }}
                    
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 25px 0;
                        font-size: 0.95rem;
                        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
                        border-radius: 8px;
                        overflow: hidden;
                    }}
                    
                    table th, table td {{
                        border: 1px solid #e2e8f0;
                        padding: 15px;
                        text-align: left;
                    }}
                    
                    table th {{
                        background-color: #667eea;
                        color: white;
                        font-weight: 600;
                        text-transform: uppercase;
                        font-size: 0.85rem;
                        letter-spacing: 0.5px;
                    }}
                    
                    table tr:nth-child(even) {{
                        background-color: #f7fafc;
                    }}
                    
                    table tr:hover {{
                        background-color: #edf2f7;
                    }}
                    
                    ul, ol {{
                        margin-bottom: 1.5rem;
                        padding-left: 2.5rem;
                        color: #2d3748;
                    }}
                    
                    li {{
                        margin-bottom: 0.75rem;
                        line-height: 1.6;
                    }}
                    
                    a {{
                        color: #667eea;
                        text-decoration: none;
                        border-bottom: 1px solid transparent;
                        transition: border-bottom 0.2s ease;
                    }}
                    
                    a:hover {{
                        border-bottom: 1px solid #667eea;
                    }}
                    
                    img {{
                        max-width: 100%;
                        height: auto;
                        border-radius: 8px;
                        margin: 15px 0;
                        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
                    }}
                    
                    hr {{
                        border: none;
                        border-top: 2px solid #e2e8f0;
                        margin: 40px 0;
                        height: 1px;
                    }}
                    
                    /* Print styles for excellent PDF output */
                    @media print {{
                        body {{
                            margin: 0;
                            padding: 30px;
                            font-size: 12pt;
                            line-height: 1.6;
                        }}
                        
                        h1, h2, h3, h4, h5, h6 {{
                            page-break-after: avoid;
                            margin-top: 20pt;
                            margin-bottom: 10pt;
                        }}
                        
                        h1 {{
                            font-size: 24pt;
                            margin-top: 0;
                        }}
                        
                        h2 {{
                            font-size: 18pt;
                        }}
                        
                        h3 {{
                            font-size: 14pt;
                        }}
                        
                        pre, blockquote, table {{
                            page-break-inside: avoid;
                            margin: 15pt 0;
                        }}
                        
                        p {{
                            margin-bottom: 12pt;
                            text-align: justify;
                        }}
                        
                        a {{
                            color: #000;
                            text-decoration: underline;
                        }}
                        
                        code {{
                            font-size: 10pt;
                        }}
                        
                        pre {{
                            font-size: 10pt;
                            padding: 15pt;
                        }}
                        
                        table {{
                            font-size: 10pt;
                        }}
                        
                        table th, table td {{
                            padding: 8pt;
                        }}
                        
                        ul, ol {{
                            margin-bottom: 12pt;
                        }}
                        
                        li {{
                            margin-bottom: 6pt;
                        }}
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Get original filename without extension and unique ID
            original_name = filename.split('_', 1)[1] if '_' in filename else filename
            base_name = os.path.splitext(original_name)[0]
            
            return Response(
                styled_html,
                mimetype='text/html',
                headers={
                    'Content-Disposition': f'attachment; filename="{base_name}.html"',
                    'Content-Type': 'text/html'
                }
            )
            
        elif format == 'docx':
            # Convert to Word document with improved formatting
            doc = Document()
            
            # Set document styles
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.line_spacing = 1.15
            
            # Process markdown content
            lines = md_content.split('\n')
            in_list = False
            list_type = None
            
            for line in lines:
                line = line.strip()
                
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], 0)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], 1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], 2)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif line.startswith('#### '):
                    heading = doc.add_heading(line[5:], 3)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif line.startswith('- ') or line.startswith('* '):
                    # Unordered list
                    if not in_list or list_type != 'unordered':
                        in_list = True
                        list_type = 'unordered'
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
                    # Ordered list
                    if not in_list or list_type != 'ordered':
                        in_list = True
                        list_type = 'ordered'
                    p = doc.add_paragraph(line[3:], style='List Number')
                elif line.startswith('```'):
                    # Code block
                    if not in_list:
                        in_list = False
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                elif line.startswith('> '):
                    # Blockquote
                    if not in_list:
                        in_list = False
                    p = doc.add_paragraph(line[2:])
                    p.style = 'Quote'
                elif line == '':
                    # Empty line - end lists
                    if in_list:
                        in_list = False
                        list_type = None
                elif line and not line.startswith('#'):
                    # Regular paragraph
                    if not in_list:
                        in_list = False
                        list_type = None
                    
                    # Handle inline formatting
                    formatted_text = line
                    formatted_text = formatted_text.replace('**', '**')  # Bold
                    formatted_text = formatted_text.replace('*', '*')    # Italic
                    formatted_text = formatted_text.replace('`', '`')    # Code
                    
                    p = doc.add_paragraph(formatted_text)
                    
                    # Apply formatting
                    for run in p.runs:
                        if '**' in run.text:
                            run.bold = True
                        if '*' in run.text and not '**' in run.text:
                            run.italic = True
                        if '`' in run.text:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
            
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            # Get original filename without extension and unique ID
            original_name = filename.split('_', 1)[1] if '_' in filename else filename
            base_name = os.path.splitext(original_name)[0]
            
            return send_file(
                doc_io,
                as_attachment=True,
                download_name=f'{base_name}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            flash('Invalid download format', 'error')
            return redirect(url_for('index'))
            
    except Exception as e:
        logging.error(f"Download error: {e}")
        flash(f'Error converting file: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/api/send-ad-inquiry', methods=['POST'])
def send_ad_inquiry():
    """Handle advertiser contact form submissions with reCAPTCHA verification"""
    try:
        # Get JSON data from request
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'message': 'No data received'}), 400
        
        # Extract form fields
        name = data.get('name', '').strip()
        email = data.get('email', '').strip()
        company = data.get('company', '').strip()
        message = data.get('message', '').strip()
        recaptcha_response = data.get('g-recaptcha-response', '')
        
        # Validate required fields
        if not name or not email or not message:
            return jsonify({'success': False, 'message': 'Please fill in all required fields'}), 400
        
        # Validate email format
        if '@' not in email or '.' not in email:
            return jsonify({'success': False, 'message': 'Please enter a valid email address'}), 400
        
        # Verify reCAPTCHA
        if not recaptcha_response:
            return jsonify({'success': False, 'message': 'Please complete the reCAPTCHA'}), 400
        
        if not verify_recaptcha(recaptcha_response):
            return jsonify({'success': False, 'message': 'reCAPTCHA verification failed. Please try again.'}), 400
        
        # Log the inquiry (in production, you'd send an email or store in database)
        inquiry_data = {
            'name': name,
            'email': email,
            'company': company,
            'message': message,
            'timestamp': datetime.now().isoformat(),
            'ip_address': request.remote_addr
        }
        
        logging.info(f"New ad inquiry received: {inquiry_data}")
        
        # In production, you would:
        # 1. Send email notification to admin
        # 2. Store in database
        # 3. Send confirmation email to user
        
        return jsonify({
            'success': True, 
            'message': 'Thank you! Your inquiry has been received. We will contact you soon.'
        }), 200
        
    except Exception as e:
        logging.error(f"Error processing ad inquiry: {e}")
        return jsonify({'success': False, 'message': 'An error occurred. Please try again.'}), 500

@app.route('/blog')
def blog_index():
    """Blog index page"""
    try:
        return render_template('blog_index.html')
    except Exception as e:
        logging.error(f"Error in blog index route: {e}")
        return "Blog coming soon!", 404

@app.route('/blog/<post_name>')
def blog_post(post_name):
    """Serve individual blog posts"""
    try:
        # Map post names to their HTML files
        post_files = {
            'convert-markdown-to-pdf-online-free': 'blog/convert-markdown-to-pdf-online-free.html',
            'best-free-markdown-to-pdf-converters-2025': 'blog/best-free-markdown-to-pdf-converters-2025.html',
            'convert-markdown-to-word-documents-online': 'blog/convert-markdown-to-word-documents-online.html',
            'markdown-tips-tricks-technical-writers': 'blog/markdown-tips-tricks-technical-writers.html',
            'troubleshooting-markdown-to-pdf-conversion': 'blog/troubleshooting-markdown-to-pdf-conversion.html'
        }
        
        if post_name in post_files:
            with open(post_files[post_name], 'r', encoding='utf-8') as f:
                content = f.read()
            return content
        else:
            return "Blog post not found", 404
    except Exception as e:
        logging.error(f"Error serving blog post {post_name}: {e}")
        return "Blog post not found", 404

@app.errorhandler(413)
def too_large(e):
    flash('File too large. Maximum size is 16MB.', 'error')
    return redirect(url_for('index'))

@app.errorhandler(404)
def not_found(e):
    return render_template('index.html'), 404

@app.errorhandler(500)
def server_error(e):
    flash('An internal server error occurred. Please try again.', 'error')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)
