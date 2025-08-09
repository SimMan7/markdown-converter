# Vercel API entry point for Flask app
import sys
import os
import json
import markdown
import uuid
import tempfile
import base64
from datetime import datetime
from werkzeug.utils import secure_filename
from io import BytesIO

# Import required libraries for PDF and DOCX generation
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# In-memory storage for Vercel (temporary)
file_storage = {}

def allowed_file(filename):
    """Check if file has allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in {'md', 'markdown'}

def generate_pdf(html_content, filename):
    """Generate PDF from HTML content"""
    if not WEASYPRINT_AVAILABLE:
        return None, "WeasyPrint not available"
    
    try:
        # Create styled HTML
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: 'Arial', sans-serif; line-height: 1.6; margin: 40px; }}
                h1 {{ color: #333; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
                h2 {{ color: #444; margin-top: 30px; }}
                h3 {{ color: #555; }}
                code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
                pre {{ background: #f8f8f8; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                blockquote {{ border-left: 4px solid #667eea; padding-left: 15px; margin: 20px 0; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        font_config = FontConfiguration()
        html_doc = HTML(string=styled_html)
        pdf = html_doc.write_pdf(
            stylesheets=[],
            font_config=font_config,
            optimize_images=False
        )
        
        return pdf, None
    except Exception as e:
        return None, str(e)

def generate_docx(md_content, filename):
    """Generate DOCX from markdown content"""
    if not DOCX_AVAILABLE:
        return None, "python-docx not available"
    
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(filename.replace('.md', '').replace('.markdown', ''), 0)
        
        # Convert markdown to HTML for parsing
        html_content = markdown.markdown(
            md_content, 
            extensions=['extra', 'tables', 'codehilite', 'toc']
        )
        
        # Simple parsing of HTML to add content
        lines = html_content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('<h1>'):
                doc.add_heading(line.replace('<h1>', '').replace('</h1>', ''), 1)
            elif line.startswith('<h2>'):
                doc.add_heading(line.replace('<h2>', '').replace('</h2>', ''), 2)
            elif line.startswith('<h3>'):
                doc.add_heading(line.replace('<h3>', '').replace('</h3>', ''), 3)
            elif line.startswith('<p>') and line.endswith('</p>'):
                text = line.replace('<p>', '').replace('</p>', '')
                doc.add_paragraph(text)
            elif line.startswith('<ul>') or line.startswith('<ol>'):
                # Handle lists
                pass
            elif line.startswith('<li>'):
                text = line.replace('<li>', '').replace('</li>', '')
                doc.add_paragraph(text, style='List Bullet')
            elif line.startswith('<blockquote>'):
                text = line.replace('<blockquote><p>', '').replace('</p></blockquote>', '')
                doc.add_paragraph(text, style='Quote')
            elif line.startswith('<pre>'):
                # Handle code blocks
                pass
            elif line and not line.startswith('<'):
                # Plain text
                doc.add_paragraph(line)
        
        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue(), None
    except Exception as e:
        return None, str(e)

def handler(request):
    """Vercel serverless function handler for all routes"""
    path = request.path
    method = request.method
    
    # Handle ads.txt
    if path == '/ads.txt':
        ads_content = "google.com, ca-pub-6865885814212781, DIRECT, f08c47fec0942fa0"
        return ads_content, 200, {
            'Content-Type': 'text/plain',
            'Cache-Control': 'public, max-age=3600'
        }
    
    # Handle debug route
    if path == '/debug':
        debug_info = {
            'cwd': os.getcwd(),
            'files_in_cwd': os.listdir('.'),
            'python_version': sys.version,
            'request_method': method,
            'request_path': path,
            'file_storage_count': len(file_storage),
            'weasyprint_available': WEASYPRINT_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
            'temp_dir': tempfile.gettempdir(),
            'environment': 'vercel'
        }
        return json.dumps(debug_info, indent=2), 200, {'Content-Type': 'application/json'}
    
    # Handle upload
    if path == '/upload' and method == 'POST':
        try:
            if 'file' not in request.files:
                return json.dumps({'error': 'No file selected'}), 400, {'Content-Type': 'application/json'}
            
            file = request.files['file']
            
            if file.filename == '':
                return json.dumps({'error': 'No file selected'}), 400, {'Content-Type': 'application/json'}
            
            if not allowed_file(file.filename):
                return json.dumps({'error': 'Please upload a valid Markdown file (.md or .markdown)'}), 400, {'Content-Type': 'application/json'}
            
            # Generate unique filename
            unique_id = str(uuid.uuid4())
            filename = f"{unique_id}_{secure_filename(file.filename or 'upload')}"
            
            # Read file content
            content = file.read().decode('utf-8')
            
            # Store in memory
            file_storage[filename] = {
                'content': content,
                'original_filename': file.filename,
                'upload_time': datetime.now().isoformat()
            }
            
            # Convert markdown to HTML for preview
            html_content = markdown.markdown(
                content, 
                extensions=['extra', 'tables', 'codehilite', 'toc']
            )
            
            return json.dumps({
                'success': True,
                'filename': filename,
                'original_filename': file.filename,
                'preview_html': html_content,
                'message': 'File uploaded successfully!'
            }), 200, {'Content-Type': 'application/json'}
            
        except Exception as e:
            return json.dumps({
                'success': False,
                'error': f'Error uploading file: {str(e)}'
            }), 500, {'Content-Type': 'application/json'}
    
    # Handle download routes
    if path.startswith('/download/'):
        try:
            parts = path.split('/')
            if len(parts) >= 4:
                format_type = parts[2]  # pdf or docx
                filename = parts[3]
                
                if filename not in file_storage:
                    return json.dumps({'error': 'File not found'}), 404, {'Content-Type': 'application/json'}
                
                file_data = file_storage[filename]
                md_content = file_data['content']
                
                if format_type == 'pdf':
                    pdf_data, error = generate_pdf(
                        markdown.markdown(md_content, extensions=['extra', 'tables', 'codehilite', 'toc']),
                        filename
                    )
                    if error:
                        return json.dumps({'error': f'PDF generation failed: {error}'}), 500, {'Content-Type': 'application/json'}
                    
                    # Return PDF as base64
                    pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
                    return json.dumps({
                        'success': True,
                        'format': 'pdf',
                        'filename': filename.replace('.md', '.pdf').replace('.markdown', '.pdf'),
                        'data': pdf_base64,
                        'message': 'PDF generated successfully'
                    }), 200, {'Content-Type': 'application/json'}
                
                elif format_type == 'docx':
                    docx_data, error = generate_docx(md_content, filename)
                    if error:
                        return json.dumps({'error': f'DOCX generation failed: {error}'}), 500, {'Content-Type': 'application/json'}
                    
                    # Return DOCX as base64
                    docx_base64 = base64.b64encode(docx_data).decode('utf-8')
                    return json.dumps({
                        'success': True,
                        'format': 'docx',
                        'filename': filename.replace('.md', '.docx').replace('.markdown', '.docx'),
                        'data': docx_base64,
                        'message': 'DOCX generated successfully'
                    }), 200, {'Content-Type': 'application/json'}
                
                else:
                    return json.dumps({'error': 'Invalid format. Use pdf or docx'}), 400, {'Content-Type': 'application/json'}
            else:
                return json.dumps({'error': 'Invalid download path'}), 400, {'Content-Type': 'application/json'}
        except Exception as e:
            return json.dumps({'error': f'Download error: {str(e)}'}), 500, {'Content-Type': 'application/json'}
    
    # Default response for other routes
    return json.dumps({
        'message': 'Markdown to PDF Converter API',
        'available_routes': ['/upload', '/download/{format}/{filename}', '/ads.txt', '/debug'],
        'status': 'running',
        'weasyprint_available': WEASYPRINT_AVAILABLE,
        'docx_available': DOCX_AVAILABLE
    }), 200, {'Content-Type': 'application/json'}
