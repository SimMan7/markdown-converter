import os
import logging
import markdown
import json
import base64
import requests
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, jsonify, Response
from werkzeug.utils import secure_filename
from docx import Document
import io
import tempfile
import uuid
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Create Flask app
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "dev-secret-key-change-in-production")

# Configuration
ALLOWED_EXTENSIONS = {'md', 'markdown'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# In-memory storage for Vercel (since filesystem is read-only)
file_storage = {}

# reCAPTCHA configuration
RECAPTCHA_SECRET_KEY = "6LeSYJ0rAAAAABx3xOqWudqBdr36gK6IcTUnBgaK"
RECAPTCHA_VERIFY_URL = "https://www.google.com/recaptcha/api/siteverify"

def allowed_file(filename):
    """Check if file has allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def verify_recaptcha(recaptcha_response):
    """Verify reCAPTCHA response with Google's API"""
    try:
        data = {
            'secret': RECAPTCHA_SECRET_KEY,
            'response': recaptcha_response
        }
        
        response = requests.post(RECAPTCHA_VERIFY_URL, data=data, timeout=10)
        result = response.json()
        
        logging.info(f"reCAPTCHA verification result: {result}")
        
        return result.get('success', False)
    except Exception as e:
        logging.error(f"Error verifying reCAPTCHA: {e}")
        return False

def convert_to_pdf_html(content):
    """Convert markdown to styled HTML for PDF conversion"""
    html_content = markdown.markdown(
        content, 
        extensions=['extra', 'tables', 'codehilite', 'toc']
    )
    
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Markdown Document</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: white;
            }}
            h1, h2, h3, h4, h5, h6 {{
                color: #2c3e50;
                margin-top: 24px;
                margin-bottom: 12px;
                page-break-after: avoid;
            }}
            h1 {{
                font-size: 2rem;
                border-bottom: 2px solid #007bff;
                padding-bottom: 0.5rem;
            }}
            h2 {{
                font-size: 1.5rem;
                border-bottom: 1px solid #e9ecef;
                padding-bottom: 0.25rem;
            }}
            p {{
                margin-bottom: 1rem;
            }}
            code {{
                background-color: #f8f9fa;
                color: #e83e8c;
                padding: 2px 4px;
                border-radius: 3px;
                font-family: 'Monaco', 'Courier New', monospace;
                font-size: 0.875em;
            }}
            pre {{
                background-color: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 5px;
                padding: 12px;
                overflow-x: auto;
                margin-bottom: 1rem;
            }}
            pre code {{
                background-color: transparent;
                padding: 0;
                color: #333;
            }}
            blockquote {{
                border-left: 4px solid #007bff;
                margin: 1rem 0;
                padding-left: 16px;
                color: #6c757d;
                font-style: italic;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 16px 0;
                font-size: 0.9rem;
            }}
            table th, table td {{
                border: 1px solid #dee2e6;
                padding: 8px;
                text-align: left;
            }}
            table th {{
                background-color: #f8f9fa;
                font-weight: 600;
            }}
            ul, ol {{
                margin-bottom: 1rem;
                padding-left: 2rem;
            }}
            li {{
                margin-bottom: 0.25rem;
            }}
            a {{
                color: #007bff;
                text-decoration: none;
            }}
            a:hover {{
                text-decoration: underline;
            }}
            img {{
                max-width: 100%;
                height: auto;
                border-radius: 4px;
                margin: 0.5rem 0;
            }}
            hr {{
                border: none;
                border-top: 2px solid #e9ecef;
                margin: 2rem 0;
            }}
            @media print {{
                body {{
                    margin: 0;
                    padding: 20px;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    page-break-after: avoid;
                }}
                pre, blockquote {{
                    page-break-inside: avoid;
                }}
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    return styled_html

@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/ads.txt')
def ads_txt():
    """Serve ads.txt file for Google AdSense"""
    try:
        return Response(
            "google.com, ca-pub-6865885814212781, DIRECT, f08c47fec0942fa0",
            mimetype='text/plain',
            headers={'Cache-Control': 'public, max-age=3600'}
        )
    except Exception as e:
        logging.error(f"Error serving ads.txt: {e}")
        return Response("google.com, ca-pub-6865885814212781, DIRECT, f08c47fec0942fa0", mimetype='text/plain')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload"""
    try:
        if 'file' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('index'))
        
        file = request.files['file']
        
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('index'))
        
        if not allowed_file(file.filename):
            flash('Please upload a valid Markdown file (.md or .markdown)', 'error')
            return redirect(url_for('index'))
        
        # Generate unique filename to avoid conflicts
        unique_id = str(uuid.uuid4())
        filename = f"{unique_id}_{secure_filename(file.filename or 'upload')}"
        
        # Read file content
        content = file.read().decode('utf-8')
        
        # Store in memory (in production, use cloud storage)
        file_storage[filename] = {
            'content': content,
            'original_filename': file.filename,
            'upload_time': datetime.now()
        }
        
        # Convert markdown to HTML for preview
        html_content = markdown.markdown(
            content, 
            extensions=['extra', 'tables', 'codehilite', 'toc']
        )
        
        flash('File uploaded successfully!', 'success')
        return render_template('index.html', 
                             preview_content=html_content, 
                             filename=filename,
                             original_filename=file.filename)
        
    except Exception as e:
        logging.error(f"Upload error: {e}")
        flash(f'Error uploading file: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/download/<format>/<filename>')
def download_file(format, filename):
    """Download converted file"""
    try:
        if filename not in file_storage:
            flash('File not found. Please upload a file first.', 'error')
            return redirect(url_for('index'))
        
        file_data = file_storage[filename]
        md_content = file_data['content']
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content, 
            extensions=['extra', 'tables', 'codehilite', 'toc']
        )
        
        if format == 'pdf':
            # For Vercel, we'll return HTML that can be printed to PDF
            # In production, use a cloud PDF service
            styled_html = convert_to_pdf_html(md_content)
            
            # Get original filename without extension and unique ID
            original_name = filename.split('_', 1)[1] if '_' in filename else filename
            base_name = os.path.splitext(original_name)[0]
            
            return Response(
                styled_html,
                mimetype='text/html',
                headers={
                    'Content-Disposition': f'attachment; filename="{base_name}.html"',
                    'Content-Type': 'text/html'
                }
            )
            
        elif format == 'docx':
            # Convert to Word document
            doc = Document()
            
            # Add title if present
            lines = md_content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], 0)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 1)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 2)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], 3)
                elif line and not line.startswith('#'):
                    doc.add_paragraph(line)
            
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            # Get original filename without extension and unique ID
            original_name = filename.split('_', 1)[1] if '_' in filename else filename
            base_name = os.path.splitext(original_name)[0]
            
            return send_file(
                doc_io,
                as_attachment=True,
                download_name=f'{base_name}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            flash('Invalid download format', 'error')
            return redirect(url_for('index'))
            
    except Exception as e:
        logging.error(f"Download error: {e}")
        flash(f'Error converting file: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/api/send-ad-inquiry', methods=['POST'])
def send_ad_inquiry():
    """Handle advertiser contact form submissions with reCAPTCHA verification"""
    try:
        # Get JSON data from request
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'message': 'No data received'}), 400
        
        # Extract form fields
        name = data.get('name', '').strip()
        email = data.get('email', '').strip()
        company = data.get('company', '').strip()
        message = data.get('message', '').strip()
        recaptcha_response = data.get('g-recaptcha-response', '')
        
        # Validate required fields
        if not name or not email or not message:
            return jsonify({'success': False, 'message': 'Please fill in all required fields'}), 400
        
        # Validate email format
        if '@' not in email or '.' not in email:
            return jsonify({'success': False, 'message': 'Please enter a valid email address'}), 400
        
        # Verify reCAPTCHA
        if not recaptcha_response:
            return jsonify({'success': False, 'message': 'Please complete the reCAPTCHA'}), 400
        
        if not verify_recaptcha(recaptcha_response):
            return jsonify({'success': False, 'message': 'reCAPTCHA verification failed. Please try again.'}), 400
        
        # Log the inquiry (in production, you'd send an email or store in database)
        inquiry_data = {
            'name': name,
            'email': email,
            'company': company,
            'message': message,
            'timestamp': datetime.now().isoformat(),
            'ip_address': request.remote_addr
        }
        
        logging.info(f"New ad inquiry received: {inquiry_data}")
        
        # In production, you would:
        # 1. Send email notification to admin
        # 2. Store in database
        # 3. Send confirmation email to user
        
        return jsonify({
            'success': True, 
            'message': 'Thank you! Your inquiry has been received. We will contact you soon.'
        }), 200
        
    except Exception as e:
        logging.error(f"Error processing ad inquiry: {e}")
        return jsonify({'success': False, 'message': 'An error occurred. Please try again.'}), 500

@app.errorhandler(413)
def too_large(e):
    flash('File too large. Maximum size is 16MB.', 'error')
    return redirect(url_for('index'))

@app.errorhandler(404)
def not_found(e):
    return render_template('index.html'), 404

@app.errorhandler(500)
def server_error(e):
    flash('An internal server error occurred. Please try again.', 'error')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True) 